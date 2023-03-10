from docx import Document
import xlrd
import os
from docx.enum.table import WD_TABLE_ALIGNMENT
# 写入表格
# d = Document()
# t = d.add_table(rows=6, cols=12, style="Table Grid")
# cell = t.cell(0,0)
# cell.text="第一行第一列"
# d.save("./docx_Res/test.docx")

# 获取表格数据
# d = Document("./docx_Res/1.docx")
# for table in d.tables:
#    for row_index, row in enumerate(table.rows):
#        for col_index, cell in enumerate(row.cells):
#            print('pos index is ({}, {})'.format(row_index, col_index))
#            print('cell text is {}'.format(cell.text))

# 获取文件名（暂时不用）
def myrename(path):
    file_name_list = os.listdir(path)
    file_name = str(file_name_list)
    file_name = file_name.replace(".xlsx", "").replace(".xls", "")
#    print(file_name)
    return file_name


path = "D:\\vscode_PythonSource\\docx_Res\\taizhang\\"
path2 = "D:\\vscode_PythonSource\\docx_Res\\"

file_name_list = os.listdir(path)
#循环读取path中台账文件
for i, file in enumerate(file_name_list):
    try:
        print(i)
        print(file)
        # 读取Excel数据
        readFile = xlrd.open_workbook(path + file)
        # print(readFile)
        # 获取Sheet
        names = readFile.sheet_names()
        # print(names)
        # 获取Sheet对象
        obj_Sheet = readFile.sheet_by_name("Sheet1")
        # 获取行数
        # row = obj_Sheet.nrows
        col = obj_Sheet.ncols
        # print("row: ", row)
        # print("col: ", col)
        # 获取行数据
        # for i in range(row):
        #    print(obj_Sheet.row_values(i))
        # 获取列数据
        # for i in range(col):
        #    print(obj_Sheet.col_values(i))
        # 获取单元格数据
        # for i in range(row):
        #    for j in range(col):
        #        print(obj_Sheet.cell_value(i, j))
        # 单元格行列均从0开始
        # print(obj_Sheet.cell_value(1,1))
        # print(obj_Sheet.row_values(1)[0])
        # print(len(obj_Sheet.row_values(1)))

        # 根据获取的行数创建Word表格
        row = obj_Sheet.nrows
        doc = Document()
        # 需要的表格为行数+4，列数13，固定
        table = doc.add_table(rows=row+3, cols=13, style="Table Grid")
        # 填入表头、表尾合计以及最左侧
        # 表头
        head_cells = table.rows[0].cells
        head_cells[0].text = '1、主机和存储资源'
        # 合并单元格,从第一列合并到最后一列
        head_cells[0].merge(head_cells[-1])
        # 设置居中显示
        head_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        head_cells = table.rows[1].cells
        # 20221123修改,去掉主机类型，加入IP地址
        # head_cells[0].text = '主机类型'
        # 纵向合并单元格
        # left_cells = table.columns[0].cells
        # left_cells[1].merge(left_cells[-1])

        head_cells[0].text = '序号'
        # 合并单元格
        left_cells = table.columns[0].cells
        head_cells[0].merge(left_cells[2])

        head_cells[1].text = '主机名称'
        # 合并单元格
        left_cells = table.columns[1].cells
        head_cells[1].merge(left_cells[2])

        # 20221123加一列，IP地址
        head_cells[2].text = 'IP'
        # 合并单元格
        left_cells = table.columns[2].cells
        head_cells[2].merge(left_cells[2])

        head_cells[3].text = 'vCPU'
        # 合并单元格
        left_cells = table.columns[3].cells
        head_cells[3].merge(left_cells[2])

        head_cells[4].text = '内存'
        # 合并单元格
        left_cells = table.columns[4].cells
        head_cells[4].merge(left_cells[2])

        head_cells[5].text = '系统磁盘'
        head_cells[6].text = ''
        # 合并单元格
        head_cells[5].merge(head_cells[6])

        head_cells[7].text = '数据磁盘'
        head_cells[8].text = ''
        # 合并单元格
        head_cells[7].merge(head_cells[8])

        head_cells[9].text = '操作系统'
        # 合并单元格
        left_cells = table.columns[9].cells
        head_cells[9].merge(left_cells[2])

        head_cells[10].text = '部署网络'
        # 合并单元格
        left_cells = table.columns[10].cells
        head_cells[10].merge(left_cells[2])

        head_cells[11].text = '对外IP地址端口'
        # 合并单元格
        left_cells = table.columns[11].cells
        head_cells[11].merge(left_cells[2])

        head_cells[12].text = '数量'
        # 合并单元格
        left_cells = table.columns[12].cells
        head_cells[12].merge(left_cells[2])

        head_cells = table.rows[2].cells
        head_cells[5].text = '容量'
        head_cells[6].text = '类型'
        head_cells[7].text = '容量'
        head_cells[8].text = '类型'
        # 表格最尾部填入总计,row+3为总行数，减1为最后一行
        bottum_cells = table.rows[row+2].cells
        bottum_cells[1].text = "总计"
        # 合并单元格
        bottum_cells[1].merge(bottum_cells[2])

        # 填入excel数据
        # 循环读取excel列填入doc table列

        # 从1 开始，主机名称
        for k in range(1, len(obj_Sheet.col_values(1))):
            print(obj_Sheet.col_values(1)[k])
        # 获取需要填入此类数据的所有列，主机名称
        left_cells = table.columns[1].cells
        # 从第2列第3行开始填入主机名称数据,结束值为数据行数+3,即上面获取excel数据行数去掉表头一行，row - 1 + 3
        for j in range(3, row+2):
            left_cells[j].text = obj_Sheet.col_values(1)[j-2]
        # 以此类推
        # 20221123加入IP数值
        left_cells = table.columns[2].cells
        for j in range(3, row+2):
            left_cells[j].text = obj_Sheet.col_values(2)[j-2]

        # CPU
        left_cells = table.columns[3].cells
        for j in range(3, row+2):
            left_cells[j].text = str(int(obj_Sheet.col_values(3)[j-2]))
        # CPU总和
        sum_CPU = 0
        for i in range(1, len(obj_Sheet.col_values(3))):
            sum_CPU += int(obj_Sheet.col_values(3)[i])
        #    print(obj_Sheet.col_values(3)[i])
        # 填入总计栏
        bottum_cells[3].text = str(sum_CPU)

        # 内存
        left_cells = table.columns[4].cells
        for j in range(3, row+2):
            left_cells[j].text = str(int(obj_Sheet.col_values(4)[j-2]))
        # MEM总和
        sum_MEM = 0
        for i in range(1, len(obj_Sheet.col_values(4))):
            sum_MEM += int(obj_Sheet.col_values(4)[i])
        # 填入总计栏
        bottum_cells[4].text = str(sum_MEM)

        # 系统盘SAS
        left_cells = table.columns[5].cells
        # 系统盘类型
        left_cells_pan = table.columns[6].cells
        for j in range(3, row+2):
            #20230310添加判断，如果为零则依次取值，直到不为零则填入
            if obj_Sheet.col_values(7)[j-2] != 0 :
                left_cells[j].text = str(int(obj_Sheet.col_values(7)[j-2]))
                left_cells_pan[j].text = 'SATA'
                #资源盘类型跟随填入
            elif obj_Sheet.col_values(8)[j-2] != 0 :
                left_cells[j].text = str(int(obj_Sheet.col_values(8)[j-2]))
                left_cells_pan[j].text = 'SAS'
            else:
                left_cells[j].text = str(int(obj_Sheet.col_values(9)[j-2]))
                left_cells_pan[j].text = 'SSD'
        # 系统盘总和
        sum_XiTong = 0
        for i in range(1, len(obj_Sheet.col_values(8))):
            #20230310,添加判断，直到取到第一个不为零为止
            if obj_Sheet.col_values(8)[i] != 0 :
                sum_XiTong += int(obj_Sheet.col_values(8)[i])
            elif obj_Sheet.col_values(7)[i] != 0 :
                sum_XiTong += int(obj_Sheet.col_values(7)[i])
            else:
                sum_XiTong += int(obj_Sheet.col_values(9)[i])
        # 填入总计栏
        bottum_cells[5].text = str(sum_XiTong)
        # 系统盘类型，默认SAS，可能需要手动修改
        #20230310注释，放入磁盘大小同步写入
#        left_cells = table.columns[6].cells
#        for j in range(3, row+2):
#            left_cells[j].text = 'SAS'

        # 数据盘SAS
        left_cells = table.columns[7].cells
        # 数据盘类型
        left_cells_pan = table.columns[8].cells
        for j in range(3, row+2):
            #20230310加入,同上系统盘
            if obj_Sheet.col_values(10)[j-2] != 0 :
                left_cells[j].text = str(int(obj_Sheet.col_values(10)[j-2]))
                left_cells_pan[j].text = 'SATA'
            elif obj_Sheet.col_values(11)[j-2] != 0:
                left_cells[j].text = str(int(obj_Sheet.col_values(11)[j-2]))
                left_cells_pan[j].text = 'SAS'
            else:
                left_cells[j].text = str(int(obj_Sheet.col_values(12)[j-2]))
                left_cells_pan[j].text = 'SSD'
        # 数据盘总和
        sum_ShuJu = 0
        for i in range(1, len(obj_Sheet.col_values(11))):
            #20230310添加判断，同上
            if obj_Sheet.col_values(11)[i] != 0 :
                sum_ShuJu += int(obj_Sheet.col_values(11)[i])
            elif obj_Sheet.col_values(10)[i] != 0 :
                sum_ShuJu += int(obj_Sheet.col_values(10)[i])
            else:
                sum_ShuJu += int(obj_Sheet.col_values(12)[i])
        # 填入总计栏
        bottum_cells[7].text = str(sum_ShuJu)
        # 数据盘类型，默认SAS，可能需要手动修改
#        left_cells = table.columns[8].cells
#        for j in range(3, row+2):
#            left_cells[j].text = 'SAS'

        # 操作系统
        left_cells = table.columns[9].cells
        for j in range(3, row+2):
            left_cells[j].text = obj_Sheet.col_values(5)[j-2]
        # 部署网络
        left_cells = table.columns[10].cells
        for j in range(3, row+2):
            left_cells[j].text = obj_Sheet.col_values(6)[j-2]
        # 数量
        left_cells = table.columns[12].cells
        for j in range(3, row+2):
            left_cells[j].text = '1'
        # 序号
        left_cells = table.columns[0].cells
        xuhao = 1
        for j in range(3, row+2):
            left_cells[j].text = str(xuhao)
            xuhao = xuhao + 1
            # 序号累计最后是数量，填入总计
            bottum_cells[12].text = str(xuhao-1)

        # 合并总计栏系统盘和数据盘数量
        sum_Total = sum_ShuJu + sum_XiTong
        bottum_cells[5].merge(bottum_cells[8]).text = str(sum_Total)
        #保存文档
        doc.save(path2 + file + ".docx")
    except Exception as e:
        print(e.__cause__)
    else:
        print("提取成功！")
